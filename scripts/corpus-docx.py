#!/usr/bin/env python3
"""
Genera il corpus DOCX di prova per il motore Track Changes.

I file sono prodotti con python-docx, che scrive OOXML conforme: sono documenti
Word veri, non file finti costruiti a mano. Servono a verificare che il motore
preservi ciò che c'è dentro — stili, liste, tabelle, note, link, immagini,
intestazioni — e non solo il testo.

    python3 scripts/corpus-docx.py tests/corpus
"""
import sys
import zlib
import struct
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor

DESTINAZIONE = Path(sys.argv[1] if len(sys.argv) > 1 else "tests/corpus")
DESTINAZIONE.mkdir(parents=True, exist_ok=True)

# Frasi con errori deliberati, per avere qualcosa da correggere davvero.
FRASI = [
    "Nel mezzo del cammin di nostra vita mi ritrovai per una selva oscura.",
    "La casa era la casa di sempre, con l'acuqa che gocciolava dal pozzo.",
    "Disse che sarebbe tornato , ma non tornò mai.",
    "Il giorno dopo pioveva; il giorno dopo ancora pioveva anche.",
    "Un uomo entrò nella stanza e l'uomo si sedette sulla sedia.",
    "Non aveva ne pane ne companatico, eppure sorrideva.",
    "Le finestre davano sul cortile, dove i bambini giocavano a palla.",
    "Qual'è il senso di tutto questo, si chiedeva ogni sera.",
]


def png_minimo() -> bytes:
    """Un PNG 1x1 valido, costruito byte per byte: nessuna dipendenza grafica."""
    def blocco(tipo: bytes, dati: bytes) -> bytes:
        contenuto = tipo + dati
        return struct.pack(">I", len(dati)) + contenuto + struct.pack(
            ">I", zlib.crc32(contenuto) & 0xFFFFFFFF
        )

    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00\xff\xff\xff")
    return (
        b"\x89PNG\r\n\x1a\n"
        + blocco(b"IHDR", ihdr)
        + blocco(b"IDAT", idat)
        + blocco(b"IEND", b"")
    )


def semplice(percorso: Path) -> None:
    """Testo scorrevole, senza fronzoli: il caso base."""
    d = Document()
    d.add_paragraph("Un racconto semplice")
    for f in FRASI:
        d.add_paragraph(f)
    d.save(percorso)


def ricco(percorso: Path) -> None:
    """Tutto ciò che il capitolato chiede di preservare, in un file solo."""
    d = Document()

    d.add_heading("Titolo dell'opera", level=0)
    d.add_heading("Capitolo primo", level=1)

    p = d.add_paragraph()
    p.add_run("Un paragrafo con ").bold = False
    p.add_run("grassetto").bold = True
    p.add_run(", ")
    p.add_run("corsivo").italic = True
    p.add_run(" e ")
    sottolineato = p.add_run("sottolineato")
    sottolineato.underline = True
    p.add_run(". La casa era la casa di sempre, con l'acuqa nel pozzo.")

    colorato = d.add_paragraph().add_run("Testo colorato e in altro corpo.")
    colorato.font.size = Pt(14)
    colorato.font.color.rgb = RGBColor(0x6C, 0x4B, 0xFF)

    giustificato = d.add_paragraph(FRASI[0])
    giustificato.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    d.add_heading("Elenchi", level=2)
    for voce in ["Primo punto", "Secondo punto , con uno spazio di troppo", "Terzo punto"]:
        d.add_paragraph(voce, style="List Bullet")
    for voce in ["Numero uno", "Numero due", "Numero tre"]:
        d.add_paragraph(voce, style="List Number")

    d.add_heading("Una tabella", level=2)
    tabella = d.add_table(rows=3, cols=3)
    tabella.style = "Table Grid"
    intestazioni = ["Capitolo", "Parole", "Note"]
    for i, testo in enumerate(intestazioni):
        tabella.cell(0, i).text = testo
    for r in range(1, 3):
        tabella.cell(r, 0).text = f"Capitolo {r}"
        tabella.cell(r, 1).text = str(r * 1200)
        tabella.cell(r, 2).text = "Da rivedere , forse"

    d.add_heading("Un'immagine", level=2)
    immagine = DESTINAZIONE / "_immagine.png"
    immagine.write_bytes(png_minimo())
    d.add_picture(str(immagine), width=Inches(1))
    d.add_paragraph("Didascalia dell'immagine.", style="Caption")

    d.add_heading("Un collegamento", level=2)
    # python-docx non ha un'API per gli hyperlink: si costruisce la relazione
    # e l'elemento a mano, che è anche il modo in cui li scrive Word.
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    paragrafo = d.add_paragraph("Vedi ")
    rel_id = d.part.relate_to(
        "https://proemios.it",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    proprieta = OxmlElement("w:rPr")
    stile = OxmlElement("w:rStyle")
    stile.set(qn("w:val"), "Hyperlink")
    proprieta.append(stile)
    run.append(proprieta)
    testo = OxmlElement("w:t")
    testo.text = "il sito dello studio"
    run.append(testo)
    link.append(run)
    paragrafo._p.append(link)
    paragrafo.add_run(" per i dettagli.")

    sezione = d.sections[0]
    sezione.header.paragraphs[0].text = "Proemios — bozza riservata"
    sezione.footer.paragraphs[0].text = "Pagina"

    d.add_page_break()
    d.add_heading("Capitolo secondo", level=1)
    for f in FRASI:
        d.add_paragraph(f)

    d.save(percorso)
    immagine.unlink(missing_ok=True)


def lungo(percorso: Path, parole_obiettivo: int = 82_000) -> None:
    """Un manoscritto della lunghezza di un romanzo vero."""
    d = Document()
    d.add_heading("Un romanzo lungo", level=0)

    parole = 0
    capitolo = 0
    # Il contatore dei paragrafi, non quello delle parole, fa ruotare le frasi:
    # il conteggio parole cresce a passi multipli della lunghezza dell'elenco e
    # lascerebbe l'indice fermo.
    paragrafo = 0
    while parole < parole_obiettivo:
        capitolo += 1
        d.add_heading(f"Capitolo {capitolo}", level=1)
        for _ in range(40):
            frase = " ".join(FRASI[(paragrafo + k) % len(FRASI)] for k in range(4))
            d.add_paragraph(frase)
            paragrafo += 1
            parole += len(frase.split())
            if parole >= parole_obiettivo:
                break

    d.save(percorso)
    return parole


semplice(DESTINAZIONE / "semplice.docx")
print("semplice.docx")

ricco(DESTINAZIONE / "ricco.docx")
print("ricco.docx")

parole = lungo(DESTINAZIONE / "lungo.docx")
print(f"lungo.docx ({parole} parole circa)")
